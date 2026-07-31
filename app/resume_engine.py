"""简历文本提取、事实约束提示词和结构化分析结果整形。"""

import io
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import BaseModel, Field
from pypdf import PdfReader

from app.agent_contracts import invoke_structured
from app.config import Settings
from app.model_gateway import create_chat_model


class ResumeParseError(ValueError):
    pass


class ResumeIssue(BaseModel):
    severity: str
    category: str
    message: str
    evidence: str
    suggestion: str


class ResumeDraftSection(BaseModel):
    title: str = Field(min_length=1, max_length=100)
    items: list[str] = Field(default_factory=list, max_length=50)


class ResumeDraft(BaseModel):
    name: str = Field(default="", max_length=100)
    headline: str = Field(default="", max_length=200)
    summary: str = Field(default="", max_length=2000)
    sections: list[ResumeDraftSection] = Field(default_factory=list)
    pending_questions: list[str] = Field(default_factory=list, max_length=30)


class ResumeAnalysisResult(BaseModel):
    scores: dict[str, float]
    keyword_matches: list[str] = Field(default_factory=list)
    keyword_gaps: list[str] = Field(default_factory=list)
    issues: list[ResumeIssue] = Field(default_factory=list)
    draft: ResumeDraft


_NUMBER_TOKEN = re.compile(
    r"(?<![\d.])(?:19|20)\d{2}(?![\d.])|"
    r"(?<![\d.])\d+(?:\.\d+)?%?(?![\d.])"
)
_PLACEHOLDER = re.compile(r"\[(?:待补充|请补充)[^\]]*\]|待补充")


def parse_resume(path: Path, content_type: str) -> str:
    try:
        if content_type == "application/pdf":
            reader = PdfReader(str(path))
            text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        elif content_type.endswith("wordprocessingml.document"):
            document = Document(str(path))
            paragraphs = [item.text for item in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(
                        " | ".join(cell.text for cell in row.cells)
                    )
            text = "\n".join(paragraphs)
        else:
            raise ResumeParseError("不支持的简历格式")
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError("简历文件损坏或无法解析") from exc
    normalized = "\n".join(
        line.strip() for line in text.splitlines() if line.strip()
    )
    if len(normalized) < 20:
        raise ResumeParseError(
            "未提取到足够文本；扫描版 PDF 暂不支持，请上传可复制文本的 PDF 或 DOCX"
        )
    return normalized[:60_000]


def analyze_resume(
    *,
    resume_text: str,
    job_description: str,
    target_role: str,
    experience_level: str,
    settings: Settings,
) -> ResumeAnalysisResult:
    model = create_chat_model(
        "resume_analysis",
        temperature=0.2,
        max_tokens=min(2000, settings.llm_max_output_tokens),
        timeout_seconds=settings.resume_analysis_timeout_seconds,
        max_retries=settings.resume_analysis_max_retries,
        settings=settings,
    )
    result = invoke_structured(
        model,
        [
            SystemMessage(
                content=(
                    "你是技术岗位简历评估专家。只输出一个JSON对象，不使用Markdown。"
                    "不得编造公司、岗位、日期、数字、技术或成果。每个问题必须引用"
                    "输入简历中的短证据。缺失事实放入pending_questions。"
                    "scores必须包含match、completeness、relevance、clarity、"
                    "impact、ats，范围0到100。JSON格式："
                    '{"scores":{"match":0,"completeness":0,"relevance":0,'
                    '"clarity":0,"impact":0,"ats":0},'
                    '"keyword_matches":[],"keyword_gaps":[],'
                    '"issues":[{"severity":"high|medium|low",'
                    '"category":"","message":"","evidence":"","suggestion":""}],'
                    '"draft":{"name":"","headline":"","summary":"",'
                    '"sections":[{"title":"","items":[]}],'
                    '"pending_questions":[]}}'
                )
            ),
            HumanMessage(
                content=(
                    f"目标岗位：{target_role or '未指定'}\n"
                    f"经验级别：{experience_level or '未指定'}\n"
                    f"岗位描述：\n{job_description or '未提供，请按目标岗位评估'}\n\n"
                    f"简历原文：\n{resume_text}"
                )
            ),
        ],
        ResumeAnalysisResult,
    )
    result.scores = {
        key: max(0.0, min(100.0, float(value)))
        for key, value in result.scores.items()
    }
    required = {
        "match",
        "completeness",
        "relevance",
        "clarity",
        "impact",
        "ats",
    }
    if not required.issubset(result.scores):
        raise ValueError("简历评估缺少必需评分")
    return result


def draft_text(draft: ResumeDraft) -> str:
    parts = [draft.name, draft.headline, draft.summary]
    for section in draft.sections:
        parts.append(section.title)
        parts.extend(section.items)
    return "\n".join(item for item in parts if item)


def find_fact_warnings(
    source_text: str,
    draft: ResumeDraft,
) -> list[dict[str, str]]:
    warnings: list[dict[str, str]] = []
    source_numbers = set(_NUMBER_TOKEN.findall(source_text))
    for token in sorted(set(_NUMBER_TOKEN.findall(draft_text(draft)))):
        if token not in source_numbers:
            warnings.append(
                {
                    "code": "unsupported_number",
                    "message": f"优化稿中的数字“{token}”没有原文证据",
                }
            )
    if _PLACEHOLDER.search(draft_text(draft)) or draft.pending_questions:
        warnings.append(
            {
                "code": "pending_information",
                "message": "优化稿仍有待补充信息，请确认或删除后再导出",
            }
        )
    return warnings


def render_docx(draft: ResumeDraft) -> bytes:
    document = Document()
    if draft.name:
        document.add_heading(draft.name, level=0)
    if draft.headline:
        document.add_paragraph(draft.headline)
    if draft.summary:
        document.add_heading("个人简介", level=1)
        document.add_paragraph(draft.summary)
    for section in draft.sections:
        document.add_heading(section.title, level=1)
        for item in section.items:
            document.add_paragraph(item, style="List Bullet")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def serialize_analysis(
    result: ResumeAnalysisResult,
    *,
    source_text: str,
) -> dict[str, str]:
    warnings = find_fact_warnings(source_text, result.draft)
    report = result.model_dump(exclude={"draft"})
    return {
        "report_json": json.dumps(report, ensure_ascii=False),
        "draft_json": result.draft.model_dump_json(),
        "warnings_json": json.dumps(warnings, ensure_ascii=False),
    }


def deserialize_json(value: Any, default: object) -> object:
    if not value:
        return default
    try:
        return json.loads(str(value))
    except (TypeError, ValueError):
        return default
