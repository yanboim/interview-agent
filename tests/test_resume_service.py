import io
from pathlib import Path

import pytest
from docx import Document

from app.application.resume_service import ResumeConflict, ResumeService
from app.config import Settings
from app.resume_engine import (
    ResumeAnalysisResult,
    ResumeDraft,
    ResumeDraftSection,
    ResumeIssue,
    find_fact_warnings,
    parse_resume,
)
from app.storage import ConversationStore
from app.user_files import (
    LocalUserFileStore,
    UnsupportedUserFile,
    UserFileTooLarge,
)


def docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_heading("张三", level=0)
    document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def analysis_result() -> ResumeAnalysisResult:
    return ResumeAnalysisResult(
        scores={
            "match": 82,
            "completeness": 80,
            "relevance": 84,
            "clarity": 78,
            "impact": 75,
            "ats": 88,
        },
        keyword_matches=["Python"],
        keyword_gaps=["可观测性"],
        issues=[
            ResumeIssue(
                severity="medium",
                category="成果",
                message="成果描述可以更具体",
                evidence="接口延迟降低30%",
                suggestion="补充基线和测量口径",
            )
        ],
        draft=ResumeDraft(
            name="张三",
            headline="高级后端工程师",
            summary="8年后端研发经验，专注Python服务。",
            sections=[
                ResumeDraftSection(
                    title="项目经历",
                    items=["优化Python接口，延迟降低30%"],
                )
            ],
        ),
    )


def make_service(tmp_path: Path):
    store = ConversationStore(tmp_path / "resume.db")
    files = LocalUserFileStore(tmp_path / "files", max_upload_bytes=2_000_000)
    jobs: list[dict[str, object]] = []

    def enqueue(job_type, payload, **kwargs):
        jobs.append({"type": job_type, "payload": payload, **kwargs})
        return "job-1"

    def analyzer(**_):
        return analysis_result()

    settings = Settings(
        _env_file=None,
        resume_feature_enabled=True,
        user_files_dir=tmp_path / "files",
        resume_max_upload_bytes=2_000_000,
        zhipu_model="test-model",
    )
    service = ResumeService(
        store,
        files,
        settings,
        enqueue=enqueue,
        analyzer=analyzer,
    )
    return service, store, files, jobs


def test_resume_service_full_lifecycle(tmp_path: Path) -> None:
    service, _, files, jobs = make_service(tmp_path)
    content = docx_bytes(
        "8年Python后端研发经验。负责接口优化，延迟降低30%。"
    )

    created = service.create(
        user_id="user-1",
        original_filename="resume.docx",
        source=io.BytesIO(content),
        job_description="招聘高级Python后端工程师",
        idempotency_key="upload-key-1",
    )

    assert created["status"] == "uploaded"
    analysis_id = created["analyses"][0]["analysis_id"]
    assert jobs[0]["type"] == "resume_analysis"
    assert jobs[0]["payload"] == {"analysis_id": analysis_id}

    assert service.process_analysis(analysis_id=analysis_id)["outcome"] == "completed"

    ready = service.get(user_id="user-1", resume_id=created["resume_id"])
    analysis = ready["analyses"][0]
    assert ready["status"] == "ready"
    assert analysis["report"]["scores"]["match"] == 82
    assert analysis["warnings"] == []

    exported, filename = service.export_docx(
        user_id="user-1",
        analysis_id=analysis_id,
    )
    assert exported.startswith(b"PK")
    assert filename == "resume-optimized.docx"

    storage_key = service.repository.get_resume(
        user_id="user-1",
        resume_id=created["resume_id"],
    )["storage_key"]
    assert files.path(storage_key).is_file()
    assert service.delete(user_id="user-1", resume_id=created["resume_id"])
    assert not (tmp_path / "files" / storage_key).exists()
    assert not service.delete(user_id="user-1", resume_id=created["resume_id"])


def test_resume_upload_idempotency_and_owner_scope(tmp_path: Path) -> None:
    service, _, _, jobs = make_service(tmp_path)
    content = docx_bytes("8年Python经验，接口延迟降低30%。")
    first = service.create(
        user_id="user-1",
        original_filename="resume.docx",
        source=io.BytesIO(content),
        job_description="Python",
        idempotency_key="upload-key-1",
    )
    replay = service.create(
        user_id="user-1",
        original_filename="resume.docx",
        source=io.BytesIO(content),
        job_description="Python",
        idempotency_key="upload-key-1",
    )

    assert replay["resume_id"] == first["resume_id"]
    assert len(jobs) == 2
    with pytest.raises(ResumeConflict):
        service.create(
            user_id="user-1",
            original_filename="resume.docx",
            source=io.BytesIO(docx_bytes("不同的简历内容，具有足够文本。")),
            job_description="Python",
            idempotency_key="upload-key-1",
        )
    assert service.list(user_id="user-2") == []


def test_draft_update_uses_optimistic_revision_and_fact_guard(
    tmp_path: Path,
) -> None:
    service, _, _, _ = make_service(tmp_path)
    created = service.create(
        user_id="user-1",
        original_filename="resume.docx",
        source=io.BytesIO(
            docx_bytes("8年Python经验，接口延迟降低30%。")
        ),
        job_description="Python",
        idempotency_key="upload-key-1",
    )
    analysis_id = created["analyses"][0]["analysis_id"]
    service.process_analysis(analysis_id=analysis_id)

    changed = analysis_result().draft.model_dump()
    changed["summary"] = "10年Python经验"
    updated = service.update_draft(
        user_id="user-1",
        analysis_id=analysis_id,
        expected_revision=1,
        draft_payload=changed,
    )
    assert updated["revision"] == 2
    assert updated["warnings"][0]["code"] == "unsupported_number"
    with pytest.raises(ResumeConflict):
        service.update_draft(
            user_id="user-1",
            analysis_id=analysis_id,
            expected_revision=1,
            draft_payload=changed,
        )
    with pytest.raises(ResumeConflict):
        service.export_docx(user_id="user-1", analysis_id=analysis_id)


def test_user_file_store_validates_type_size_and_path(tmp_path: Path) -> None:
    files = LocalUserFileStore(tmp_path / "files", max_upload_bytes=100)
    with pytest.raises(UnsupportedUserFile):
        files.save(
            user_id="user-1",
            asset_id="asset-1",
            original_filename="../resume.exe",
            source=io.BytesIO(b"MZ"),
        )
    with pytest.raises(UserFileTooLarge):
        files.save(
            user_id="user-1",
            asset_id="asset-2",
            original_filename="resume.pdf",
            source=io.BytesIO(b"%PDF-" + b"x" * 100),
        )
    with pytest.raises(UnsupportedUserFile):
        files.save(
            user_id="user-1",
            asset_id="asset-3",
            original_filename="resume.pdf",
            source=io.BytesIO(b"not a pdf"),
        )


def test_fact_warning_and_scanned_pdf_error(tmp_path: Path) -> None:
    draft = ResumeDraft(
        summary="10年经验",
        pending_questions=["请补充指标"],
    )
    warnings = find_fact_warnings("8年经验", draft)
    assert {item["code"] for item in warnings} == {
        "unsupported_number",
        "pending_information",
    }
    path = tmp_path / "empty.pdf"
    path.write_bytes(
        b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\n"
        b"trailer<</Root 1 0 R>>\n%%EOF"
    )
    with pytest.raises(Exception):
        parse_resume(path, "application/pdf")
